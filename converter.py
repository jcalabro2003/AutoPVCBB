"""
converter.py - Module principal de conversion DocX vers LaTeX
PATCH: Solution robuste pour macOS read-only filesystem
"""

from docx import Document
from pathlib import Path
import re
import logging
import subprocess
import os
import sys
import tempfile
from typing import List, Optional, Tuple

from utils import TextProcessor, DocumentParser, TableProcessor
from latex_generator import LaTeXGenerator
from text_corrector import TextCorrector

logger = logging.getLogger(__name__)

def get_writable_output_dir():
    """
    Retourne un dossier accessible en écriture de manière robuste.
    Teste plusieurs emplacements et retourne le premier qui fonctionne.
    """
    # Liste des emplacements à tester, par ordre de préférence
    possible_dirs = []
    
    if sys.platform == 'darwin':  # macOS
        home = Path.home()
        possible_dirs = [
            home / "Documents" / "ConvertisseurDocxLatex",
            home / "Desktop" / "ConvertisseurDocxLatex",
            home / "Downloads" / "ConvertisseurDocxLatex",
            Path(tempfile.gettempdir()) / "ConvertisseurDocxLatex",
        ]
    elif sys.platform == 'win32':  # Windows
        if getattr(sys, 'frozen', False):
            # Exécutable : à côté de l'exe
            possible_dirs = [Path(sys.executable).parent]
        else:
            # Script : à côté du script
            possible_dirs = [Path(__file__).parent.resolve()]
        
        # Ajouter aussi Documents comme fallback
        home = Path.home()
        possible_dirs.append(home / "Documents" / "ConvertisseurDocxLatex")
    else:  # Linux
        home = Path.home()
        possible_dirs = [
            Path(__file__).parent.resolve(),
            home / "Documents" / "ConvertisseurDocxLatex",
            Path(tempfile.gettempdir()) / "ConvertisseurDocxLatex",
        ]
    
    # Tester chaque emplacement
    for directory in possible_dirs:
        try:
            # Tenter de créer le dossier
            directory.mkdir(parents=True, exist_ok=True)
            
            # Tester l'écriture avec un fichier temporaire
            test_file = directory / ".write_test"
            test_file.write_text("test")
            test_file.unlink()
            
            # Si on arrive ici, le dossier est accessible en écriture
            logger.info(f"Dossier de sortie accessible: {directory}")
            return directory
            
        except (OSError, PermissionError) as e:
            logger.debug(f" Dossier non accessible: {directory} - {e}")
            continue
    
    # Si aucun dossier ne fonctionne, utiliser le dossier temporaire système
    fallback = Path(tempfile.gettempdir()) / "ConvertisseurDocxLatex"
    logger.warning(f" Utilisation du dossier temporaire: {fallback}")
    fallback.mkdir(parents=True, exist_ok=True)
    return fallback


class DocxToLatexConverter:
    """Convertisseur principal de documents DocX vers LaTeX."""
    
    def __init__(self, config):
        self.config = config
        self.text_processor = TextProcessor(config)
        self.latex_generator = LaTeXGenerator(config, self.text_processor)
        self.text_corrector = TextCorrector(config)
        self.table_processor = TableProcessor()
        
        # Déterminer le dossier de sortie une seule fois à l'initialisation
        self.output_base_dir = get_writable_output_dir()
        logger.info(f" Dossier de sortie configuré: {self.output_base_dir}")
    
    def convert(self, docx_path: str, latex_path: str, compile_pdf: bool = True) -> Optional[str]:
        """
        Convertit un fichier DocX en fichier LaTeX.
        
        Args:
            docx_path: Chemin du fichier DocX source
            latex_path: Chemin du fichier LaTeX de sortie (ignoré, on utilise output_base_dir)
            compile_pdf: Si True, compile également le LaTeX en PDF

        Returns:
            Chemin du PDF ou du .tex généré
        """
        try:
            logger.info(f" Début de la conversion: {Path(docx_path).name}")
            
            # Charger le document
            doc = Document(docx_path)

            # Extraire le titre du document
            doc_title = Path(docx_path).stem

            # Préparer dossiers de sortie
            latex_dir = self.output_base_dir / "LaTeX"
            latex_dir.mkdir(parents=True, exist_ok=True)
            
            images_dir = latex_dir / "images"
            images_dir.mkdir(parents=True, exist_ok=True)

            logger.info(f" Fichiers LaTeX → {latex_dir}")

            # Extraire et sauvegarder les images du document
            image_map = self._extract_images(doc, images_dir)

            # Traiter le document
            document_data = self._process_document(doc, doc_title, image_map)

            # Nom du fichier .tex
            tex_filename = f"{doc_title}.tex"
            final_latex_path = latex_dir / tex_filename

            self._write_latex_file(str(final_latex_path), document_data)

            logger.info(f" Fichier LaTeX créé: {tex_filename}")

            # Compiler en PDF si demandé
            pdf_result = None
            if compile_pdf:
                pdf_path = self._compile_to_pdf(str(final_latex_path))
                if pdf_path:
                    # Créer dossier PDF
                    pdf_dir = self.output_base_dir / "PDF"
                    pdf_dir.mkdir(parents=True, exist_ok=True)

                    pdf_file = Path(pdf_path)
                    target_pdf_path = pdf_dir / pdf_file.name
                    
                    try:
                        # Tenter de déplacer
                        pdf_file.replace(target_pdf_path)
                        pdf_result = str(target_pdf_path)
                        logger.info(f" PDF créé: {pdf_file.name}")
                    except Exception as e:
                        # Fallback: copier
                        try:
                            import shutil
                            shutil.copy2(str(pdf_file), str(target_pdf_path))
                            pdf_file.unlink()
                            pdf_result = str(target_pdf_path)
                            logger.info(f" PDF créé: {pdf_file.name}")
                        except Exception as e2:
                            logger.warning(f" Impossible de déplacer le PDF: {e2}")
                            pdf_result = str(pdf_path)
                else:
                    logger.warning(" La compilation PDF a échoué")

            # Afficher un message de succès avec l'emplacement
            if pdf_result:
                logger.info(f" PDF disponible dans: {self.output_base_dir / 'PDF'}")
                return pdf_result
            else:
                logger.info(f" LaTeX disponible dans: {latex_dir}")
                return str(final_latex_path)
                
        except Exception as e:
            logger.error(f" Erreur lors de la conversion: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            raise

    def _compile_to_pdf(self, latex_path: str) -> Optional[str]:
        """
        Compile le fichier LaTeX en PDF.
        
        Args:
            latex_path: Chemin du fichier LaTeX
            
        Returns:
            Chemin du fichier PDF généré ou None si erreur
        """
        try:
            latex_file = Path(latex_path)
            pdf_file = latex_file.with_suffix('.pdf')
            
            # Dossier de travail (dossier du fichier LaTeX)
            working_dir = latex_file.parent
            
            # Nom du fichier sans extension
            file_name = latex_file.stem
            
            logger.info(f"Compilation PDF: {latex_file}")


            # Première compilation
            result = subprocess.run(
                ['pdflatex', '-interaction=nonstopmode', latex_file.name],
                cwd=working_dir,
                capture_output=True,
                text=True
            )

            combined_output = (result.stdout or '') + '\n' + (result.stderr or '')

            # If failed, try lualatex as fallback
            if result.returncode != 0:
                logger.warning("pdflatex failed, trying lualatex")
                result = subprocess.run(
                    ['lualatex', '-interaction=nonstopmode', latex_file.name],
                    cwd=working_dir,
                    capture_output=True,
                    text=True
                )
                combined_output += '\n' + (result.stdout or '') + '\n' + (result.stderr or '')

            # If still failed, analyze log for missing packages and retry once
            if result.returncode != 0:
                # search for missing .sty files
                missing = set(re.findall(r"File `([^`]+)\.sty' not found", combined_output))
                if missing:

                    # Try to add \usepackage entries for missing packages (best-effort)
                    added = []
                    for pkg in missing:
                        pkg_name = pkg.split('/')[-1]
                        usepkg = f"\\usepackage{{{pkg_name}}}"
                        if usepkg not in self.config.LATEX_PACKAGES:
                            self.config.LATEX_PACKAGES.append(usepkg)
                            added.append(pkg_name)

                    if added:
                        logger.info(f"Added missing packages to LATEX_PACKAGES: {added}. Retrying compilation.")

                        # Retry pdflatex twice
                        retry1 = subprocess.run(
                            ['pdflatex', '-interaction=nonstopmode', latex_file.name],
                            cwd=working_dir,
                            capture_output=True,
                            text=True
                        )

                        retry2 = subprocess.run(
                            ['pdflatex', '-interaction=nonstopmode', latex_file.name],
                            cwd=working_dir,
                            capture_output=True,
                            text=True
                        )

                        if retry2.returncode != 0:
                            return None
                    else:
                        return None

            # Deuxième compilation pour les références croisées (si nécessaire)
            subprocess.run(
                ['pdflatex', '-interaction=nonstopmode', latex_file.name],
                cwd=working_dir,
                capture_output=True,
                text=True
            )

            # Nettoyer les fichiers auxiliaires
            self._clean_auxiliary_files(working_dir, file_name)

            if pdf_file.exists():
                return str(pdf_file)
            else:
                return None
                
        except FileNotFoundError:
            logger.error("pdflatex n'est pas installé. Installez une distribution LaTeX (TeX Live, MiKTeX, etc.)")
            logger.error("  Pour installer LaTeX :")
            logger.error("  - Windows : Installer MiKTeX depuis https://miktex.org/")
            logger.error("  - Mac : Installer MacTeX depuis https://www.tug.org/mactex/")
            logger.error("  - Linux : sudo apt-get install texlive-full")
            return None
        except Exception as e:
            logger.error(f"Erreur lors de la compilation PDF: {e}")
            return None
    
    def _clean_auxiliary_files(self, working_dir: Path, file_name: str):
        """
        Nettoie les fichiers auxiliaires générés par LaTeX.
        
        Args:
            working_dir: Dossier de travail
            file_name: Nom du fichier sans extension
        """
        extensions_to_remove = ['.aux', '.log', '.out', '.toc', '.fls', '.fdb_latexmk']

        for ext in extensions_to_remove:
            aux_file = working_dir / f"{file_name}{ext}"
            if aux_file.exists():
                try:
                    aux_file.unlink()
                except Exception as e:
                    logger.debug(f"Impossible de supprimer {aux_file}: {e}")

    def _extract_images(self, doc: Document, images_dir: Path) -> dict:
        """
        Extrait les images présentes dans le document et les sauvegarde dans `images_dir`.

        Retourne une map {rel_id: saved_path}
        """
        image_map = {}
        try:
            for rel_id, rel in doc.part.rels.items():
                try:
                    if 'image' in rel.reltype.lower():
                        image_part = rel.target_part
                        blob = getattr(image_part, 'blob', None)
                        if not blob:
                            continue
                        filename = Path(image_part.partname).name
                        target = images_dir / filename
                        with open(target, 'wb') as f:
                            f.write(blob)
                        image_map[rel_id] = str(target)
                except Exception:
                    continue
        except Exception:
            logger.debug("Impossible d'extraire les images du document")

        return image_map

    def _emu_to_cm(self, emu: int) -> float:
        """Convertit une valeur EMU (English Metric Unit) en centimètres."""
        try:
            return (float(emu) / 914400.0) * 2.54
        except Exception:
            return 0.0
    
    def _process_document(self, doc: Document, doc_title: str, image_map: Optional[dict] = None) -> dict:
        """
        Traite le document DocX et extrait toutes les données nécessaires.
        
        Args:
            doc: Document DocX
            doc_title: Titre du document
        
        Returns:
            Dictionnaire contenant toutes les données traitées
        """
        data = {
            'title': doc_title,
            'elements': [],
            'present_names': [],
            'sections_list': [],
            'subsections_list': [],
            'first_text': None,
            'tables': list(doc.tables)
        }
        
        # Extraire la liste des sections
        data['sections_list'], data['subsections_list'] = DocumentParser.extract_sections_list(doc.paragraphs)
        
        # Collecter les paragraphes pour correction
        paragraphs_to_correct = []
        
        # Variables de suivi d'état
        table_index = 0
        in_present_section = False
        present_names_buffer = []
        
        # Parcourir les éléments du document
        for element in doc.element.body:
            element_data = self._process_element(
                element, doc, data, 
                table_index, in_present_section, 
                present_names_buffer, paragraphs_to_correct,
                image_map or {}
            )
            
            if element_data:
                if 'table_index' in element_data:
                    table_index = element_data['table_index']
                if 'in_present_section' in element_data:
                    in_present_section = element_data['in_present_section']
                if 'present_names' in element_data:
                    data['present_names'] = element_data['present_names']
        
        # Corriger tous les paragraphes en batch
        if paragraphs_to_correct:
            logger.info(f"Correction de {len(paragraphs_to_correct)} paragraphes...")
            self.text_corrector.correct_paragraphs_batch(paragraphs_to_correct)
        
        return data
    
    def _process_element(
        self, element, doc: Document, data: dict,
        table_index: int, in_present_section: bool,
        present_names_buffer: List[str],
        paragraphs_to_correct: List,
        image_map: dict
    ) -> Optional[dict]:
        """
        Traite un élément individuel du document.
        
        Returns:
            Dictionnaire avec les mises à jour d'état ou None
        """
        result = {}
        
        if element.tag.endswith("tbl"):
            # Traiter un tableau
            if table_index < len(data['tables']):
                table = data['tables'][table_index]
                result['table_index'] = table_index + 1
                
                table_data = self.table_processor.extract_table_data(
                    table, self.text_processor
                )
                table_data = self.table_processor.remove_duplicate_columns(table_data)
                
                data['elements'].append({
                    'type': 'table',
                    'data': table_data
                })
        
        elif element.tag.endswith("p"):
            # Détecter les images intégrées via r:embed dans l'XML de l'élément
            try:
                xml = element.xml
                rids = re.findall(r'r:embed="(rId[0-9]+)"', xml)
                # extraire listes cx et cy (EMU)
                cx_list = re.findall(r'cx="(\d+)"', xml)
                cy_list = re.findall(r'cy="(\d+)"', xml)

                for idx, rid in enumerate(rids):
                    if rid in image_map:
                        img_entry = {'type': 'image', 'path': image_map[rid]}
                        try:
                            cx = int(cx_list[idx]) if idx < len(cx_list) else None
                            cy = int(cy_list[idx]) if idx < len(cy_list) else None
                        except Exception:
                            cx = cy = None

                        if cx and cy:
                            # convertir EMU -> cm
                            img_entry['width_cm'] = self._emu_to_cm(cx)
                            img_entry['height_cm'] = self._emu_to_cm(cy)

                        data['elements'].append(img_entry)
            except Exception:
                pass

            # Traiter un paragraphe
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if not para:
                return None
            
            text = para.text.strip()
            
            # Premier texte du document
            if data['first_text'] is None and text:
                data['first_text'] = text
                data['elements'].append({
                    'type': 'title',
                    'text': text
                })
                return result
            
            # Gestion de la section "Présents"
            if text.startswith("Présents"):
                result['in_present_section'] = True
                return result
            
            if in_present_section:
                if not text or text.startswith("__"):
                    # Fin de la section présents
                    result['in_present_section'] = False
                    result['present_names'] = present_names_buffer.copy()
                    
                    data['elements'].append({
                        'type': 'present_section',
                        'names': present_names_buffer.copy()
                    })
                    
                    if text.startswith("__"):
                        data['elements'].append({
                            'type': 'start_text',
                            'text': text
                        })
                    
                    present_names_buffer.clear()
                else:
                    # Accumulation des noms
                    present_names_buffer.append(text)
                return result
            
            # Paragraphe normal
            if text:
                paragraphs_to_correct.append(para)
                
                if DocumentParser.is_section_header(text):
                    section_title = DocumentParser.extract_section_title(text)
                    data['elements'].append({
                        'type': 'section',
                        'title': section_title,
                        'paragraph': para
                    })
                elif DocumentParser.is_subsection_header(text):
                    subsection_title = DocumentParser.extract_section_title(text)
                    data['elements'].append({
                        'type': 'subsection',
                        'title': subsection_title,
                        'paragraph': para
                    })
                else:
                    data['elements'].append({
                        'type': 'paragraph',
                        'paragraph': para
                    })
        
        return result
    
    def _write_latex_file(self, latex_path: str, data: dict):
        """
        Écrit le fichier LaTeX à partir des données traitées.
        
        Args:
            latex_path: Chemin du fichier LaTeX de sortie
            data: Données traitées du document
        """
        with open(latex_path, 'w', encoding='utf-8') as latex_file:
            # En-tête du document
            latex_file.write(self.latex_generator.generate_document_header())
            latex_file.write(self.latex_generator.generate_title_header(data['title']) + "\n\n")
            
            # Variable pour suivre si on a déjà inséré la TOC et le logo
            toc_inserted = False
            present_inserted = False 
            
            # Traiter les éléments
            for element in data['elements']:
                elem_type = element['type']
                
                if elem_type == 'title':
                    latex_file.write(
                        self.latex_generator.generate_title_section(element['text'])
                    )
                
                elif elem_type == 'table':
                    latex_file.write(
                        self.latex_generator.generate_table(element['data'])
                    )
                elif elem_type == 'image':
                    img_path = element.get('path')
                    if img_path:
                        # Chemin relatif (images/filename) par rapport au .tex
                        img_name = Path(img_path).name
                        rel_path = Path('images') / img_name
                        width_cm = element.get('width_cm')
                        height_cm = element.get('height_cm')

                        opts = []
                        if width_cm and width_cm > 0:
                            opts.append(f"width={width_cm:.2f}cm")
                        if height_cm and height_cm > 0:
                            opts.append(f"height={height_cm:.2f}cm")

                        opt_str = f"[{','.join(opts)}]" if opts else ""

                        latex_file.write("\\begin{figure}[h]\n\\centering\n")
                        latex_file.write(f"\\includegraphics{opt_str}{{{rel_path.as_posix()}}}\n")
                        latex_file.write("\\end{figure}\n\n")
                
                elif elem_type == 'present_section':
                    names = [name.strip() for name in element['names'] if name.strip()]
                    # Filtrer les noms en enlevant les # au début
                    cleaned_names = []
                    for name in names:
                        if name.startswith("#"):
                            name = name[1:].strip()
                        cleaned_names.append(name)
                    
                    latex_file.write(
                        self.latex_generator.generate_present_section(cleaned_names)
                    )
                    present_inserted = True
                
                elif elem_type == 'start_text' or present_inserted:
                    # Insérer la TOC et le logo avant le texte de début
                    if not toc_inserted and data['sections_list']:
                        latex_file.write(self.latex_generator.generate_toc(data['sections_list'], data['subsections_list']))
                        latex_file.write(self.latex_generator.generate_logo_section())
                        toc_inserted = True
                    

                    present_inserted = False  # Reset pour éviter de réinsérer
                
                elif elem_type == 'section':
                    para = element['paragraph']
                    latex_file.write(self.latex_generator.generate_section(element['title']))

                elif elem_type == 'subsection':
                    para = element['paragraph']
                    latex_file.write(self.latex_generator.generate_subsection(element['title']))
                
                elif elem_type == 'paragraph':
                    para = element['paragraph']
        
                    latex_file.write(self._format_paragraph_with_runs(para))

            
            # Pied de page du document
            latex_file.write(self.latex_generator.generate_document_footer())
    
    def _format_paragraph_with_runs(self, para, ndlr=True) -> str:
        """
        Formate un paragraphe en conservant le formatage des runs.
        
        Args:
            para: Objet paragraphe
        
        Returns:
            Texte LaTeX formaté
        """
        latex_text = ""
        text = para.text.strip()
        fail = False

        if ":" in text:
            parts = para.text.split(":", 1)
            if len(parts) == 2:
                brut_bold_part = parts[0].strip()
                bold_part = f"\\textbf{{{self.text_processor.escape_latex(brut_bold_part)}}}"
                for i in range(len(para.runs)):
                    type = []
                    run = para.runs[i]
                    if i == 0:
                        type.append("begin")
                    if i == len(para.runs) - 1:
                        type.append("end")

                    remaining = run.text.strip()
                    if brut_bold_part in remaining:
                        remaining = remaining.replace(brut_bold_part, "")
                        remaining = remaining.replace(":", "", 1).strip()
                        
                    if remaining.startswith(":"):
                        remaining = remaining[1:].strip()

                    remaining = self.text_processor.escape_latex(remaining)
                    remaining = self.text_processor.replace_abbreviations(remaining, type=type)


                    if run.bold:
                        remaining = f"\\textbf{{{remaining}}}"
                    if run.italic:
                        remaining = f"\\textit{{{remaining}}}"
                        bold_part = f"\\textit{{{bold_part}}}"
                    
                    if remaining:
                        latex_text += remaining

                latex_text = latex_text.replace("\\textit{}", "")
                return f"{bold_part} : {latex_text.strip()}\n\n"
            else:
                fail = True
        else:
            fail = True
        

        if fail:
            for run in para.runs:
            
                run_text = run.text.strip()

                run_text = self.text_processor.escape_latex(run_text)
                run_text = self.text_processor.replace_abbreviations(run_text)
                
                if run.bold:
                    run_text = f"\\textbf{{{run_text}}}"
                if run.italic:
                    run_text = f"\\textit{{{run_text}}}"
                
                if run_text:
                    latex_text += run_text + " "
                    
            
        return latex_text.strip() + "\n\n" if latex_text.strip() else ""